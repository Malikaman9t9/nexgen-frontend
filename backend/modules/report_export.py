import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime
import io
import requests
from urllib.parse import urljoin, urlparse

def create_gauge_image(score):
    color = "#ef4444" if score < 50 else "#f59e0b" if score < 90 else "#10b981"
    fig, ax = plt.subplots(figsize=(1.5, 1.5))
    ax.pie([score, max(100-score, 0)], colors=[color, '#f1f5f9'], startangle=90, counterclock=False, wedgeprops=dict(width=0.25))
    ax.text(0, 0, str(score), ha='center', va='center', fontsize=18, fontweight='bold', color='#0f172a', family='sans-serif')
    plt.axis('off')
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', transparent=True, bbox_inches='tight', dpi=150)
    plt.close(fig)
    img_stream.seek(0)
    return img_stream

def create_pie_chart(labels, sizes, title=''):
    fig, ax = plt.subplots(figsize=(4, 3))
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=['#6366f1', '#8b5cf6', '#a78bfa', '#c4b5fd', '#ddd6fe'])
    ax.set_title(title)
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', transparent=True, bbox_inches='tight', dpi=150)
    plt.close(fig)
    img_stream.seek(0)
    return img_stream

def fetch_logo_url(url):
    try:
        # Try common logo locations
        logo_urls = [
            urljoin(url, '/logo.png'),
            urljoin(url, '/logo.svg'),
            urljoin(url, '/assets/logo.png'),
            urljoin(url, '/images/logo.png')
        ]
        for logo_url in logo_urls:
            try:
                resp = requests.get(logo_url, timeout=3)
                if resp.status_code == 200 and resp.headers.get('content-type', '').startswith('image'):
                    return logo_url
            except:
                continue
        return None
    except:
        return None

def generate_word_report(url, onpage_data, speed_data, traffic_data, ai_suggestions, agency_name="SEO Agency", client_name="Client", author_name="SEO Expert"):
    doc = Document()
    doc.sections[0].page_width = Inches(8.5)
    doc.sections[0].page_height = Inches(11)
    
    PRIMARY = RGBColor(109, 40, 217)
    SUCCESS = RGBColor(16, 185, 129)
    WARNING = RGBColor(245, 158, 11)
    DANGER = RGBColor(239, 68, 68)
    SLATE = RGBColor(100, 116, 139)
    
    # COVER
    title = doc.add_heading(f'{agency_name} - SEO Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = PRIMARY
    
    sub = doc.add_heading(f'Prepared for: {client_name}', level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = SLATE
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Target URL: {url}\n").bold = True
    meta.add_run(f"Prepared by: {author_name}\n")
    meta.add_run(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
    
    # Try to add logo if found
    logo_url = fetch_logo_url(url)
    if logo_url:
        try:
            logo_resp = requests.get(logo_url, timeout=5)
            if logo_resp.status_code == 200:
                logo_stream = io.BytesIO(logo_resp.content)
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.add_run().add_picture(logo_stream, width=Inches(2.0))
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass  # Skip logo if fetch fails
    
    doc.add_page_break()
    
    # EXECUTIVE OVERVIEW
    doc.add_heading('Executive Overview', level=1)
    overview_text = f"This report provides a comprehensive analysis of {url}'s SEO performance, covering technical on-page factors, Core Web Vitals, traffic analytics, and strategic recommendations. "
    if speed_data:
        mobile_score = speed_data.get('mobile', {}).get('performance', 0)
        desktop_score = speed_data.get('desktop', {}).get('performance', 0)
        avg_score = (mobile_score + desktop_score) // 2
        overview_text += f"The site achieves an average performance score of {avg_score}/100. "
    if traffic_data:
        visits = traffic_data.get('monthly_visits', 'N/A')
        overview_text += f"Estimated monthly visits: {visits}."
    p = doc.add_paragraph(overview_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # PERFORMANCE SCORES WITH GAUGES
    doc.add_heading('Performance Scores', level=1)
    
    if speed_data:
        # Mobile Gauge
        mobile_score = speed_data.get('mobile', {}).get('performance', 0)
        p_mobile = doc.add_paragraph()
        p_mobile.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_mobile.add_run('Mobile Performance: ').bold = True
        p_mobile.add_run(f"{mobile_score}/100")
        if mobile_score >= 90:
            p_mobile.runs[1].font.color.rgb = SUCCESS
        elif mobile_score >= 50:
            p_mobile.runs[1].font.color.rgb = WARNING
        else:
            p_mobile.runs[1].font.color.rgb = DANGER
        mobile_img = create_gauge_image(mobile_score)
        last_para = doc.paragraphs[-1]
        last_para.add_run().add_picture(mobile_img, width=Inches(1.25))
        
        # Desktop Gauge
        doc.add_paragraph()  # Space
        desktop_score = speed_data.get('desktop', {}).get('performance', 0)
        p_desktop = doc.add_paragraph()
        p_desktop.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_desktop.add_run('Desktop Performance: ').bold = True
        p_desktop.add_run(f"{desktop_score}/100")
        if desktop_score >= 90:
            p_desktop.runs[1].font.color.rgb = SUCCESS
        elif desktop_score >= 50:
            p_desktop.runs[1].font.color.rgb = WARNING
        else:
            p_desktop.runs[1].font.color.rgb = DANGER
        desktop_img = create_gauge_image(desktop_score)
        last_para = doc.paragraphs[-1]
        last_para.add_run().add_picture(desktop_img, width=Inches(1.25))
        
        # Overall Score
        doc.add_paragraph()
        overall_score = (mobile_score + desktop_score) // 2
        p_overall = doc.add_paragraph()
        p_overall.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_overall.add_run('Overall Performance Score: ').bold = True
        p_overall.add_run(f"{overall_score}/100")
        if overall_score >= 90:
            p_overall.runs[1].font.color.rgb = SUCCESS
        elif overall_score >= 50:
            p_overall.runs[1].font.color.rgb = WARNING
        else:
            p_overall.runs[1].font.color.rgb = DANGER
        p_overall.runs[0].font.size = Pt(14)
        p_overall.runs[1].font.size = Pt(14)
    
    doc.add_page_break()
    
    # ON-PAGE SEO ANALYSIS
    doc.add_heading('On-Page SEO Analysis', level=1)
    
    checks = [
        ('Title Tag', bool(onpage_data.get('title'))),
        ('Meta Description', bool(onpage_data.get('description'))),
        ('H1 Headings', len(onpage_data.get('h1', [])) > 0),
        ('Word Count >= 300', onpage_data.get('word_count', 0) >= 300),
        ('HTTPS/SSL', onpage_data.get('is_https', '') == 'Yes'),
        ('Schema Markup', bool(onpage_data.get('schema'))),
        ('Canonical URL', bool(onpage_data.get('canonical'))),
        ('No Missing ALTs', onpage_data.get('missing_alt', 0) == 0),
    ]
    
    # Create a simple table for checks
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Check'
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    hdr_cells[1].text = 'Status'
    hdr_cells[1].paragraphs[0].runs[0].bold = True
    
    for label, passed in checks:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        status_text = 'PASS' if passed else 'FAIL'
        row_cells[1].text = status_text
        if passed:
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = SUCCESS
        else:
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = DANGER
    
    doc.add_paragraph()  # Space
    
    # Additional notes
    notes = doc.add_paragraph()
    notes.add_run('Key Findings: ').bold = True
    issues = []
    if not onpage_data.get('title'):
        issues.append("Missing title tag")
    if not onpage_data.get('description'):
        issues.append("Missing meta description")
    if len(onpage_data.get('h1', [])) == 0:
        issues.append("Missing H1 heading")
    if onpage_data.get('word_count', 0) < 300:
        issues.append("Low word count (<300 words)")
    if onpage_data.get('missing_alt', 0) > 0:
        issues.append(f"{onpage_data.get('missing_alt')} images missing alt text")
    if issues:
        notes.add_run("; ".join(issues))
    else:
        notes.add_run("No critical on-page issues detected.")
    
    doc.add_page_break()
    
    # CORE WEB VITALS
    doc.add_heading('Core Web Vitals', level=1)
    
    if speed_data:
        for device in ['mobile', 'desktop']:
            device_data = speed_data.get(device, {})
            score = device_data.get('performance', 0)
            hp = doc.add_heading(f'{device.upper()} Performance: {score}/100', level=2)
            if score >= 90:
                hp.runs[0].font.color.rgb = SUCCESS
            elif score >= 50:
                hp.runs[0].font.color.rgb = WARNING
            else:
                hp.runs[0].font.color.rgb = DANGER
            
            raw_metrics = device_data.get('metrics', {})
            metrics_table = doc.add_table(rows=1, cols=2)
            metrics_table.style = 'Table Grid'
            m_hdr_cells = metrics_table.rows[0].cells
            m_hdr_cells[0].text = 'Metric'
            m_hdr_cells[1].text = 'Value'
            m_hdr_cells[0].paragraphs[0].runs[0].bold = True
            m_hdr_cells[1].paragraphs[0].runs[0].bold = True
            
            for key, label in [('fcp', 'FCP'), ('lcp', 'LCP'), ('tbt', 'TBT'), ('cls', 'CLS'), ('si', 'SI')]:
                m_row = metrics_table.add_row().cells
                value = raw_metrics.get(key, {}).get('value', 'N/A')
                m_row[0].text = label
                m_row[1].text = str(value)
    
    doc.add_page_break()
    
    # TRAFFIC ANALYTICS
    doc.add_heading('Traffic Analytics', level=1)
    
    if traffic_data:
        # Traffic Sources Pie Chart
        top_referrals = traffic_data.get('top_referrals', [])
        if top_referrals and len(top_referrals) > 0:
            labels = [item[0] for item in top_referrals[:5]]  # Top 5
            sizes = [item[1] for item in top_referrals[:5]]
            # Add "Other" for remaining traffic
            other = 100 - sum(sizes)
            if other > 0:
                labels.append('Other')
                sizes.append(other)
            pie_img = create_pie_chart(labels, sizes, 'Traffic Sources Distribution')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(pie_img, width=Inches(4.0))
        
        doc.add_paragraph()  # Space
        
        # Key Traffic Metrics
        metrics = doc.add_paragraph()
        metrics.add_run('Key Metrics: ').bold = True
        metrics.add_run(f"\n• Monthly Visits: {traffic_data.get('monthly_visits', 'N/A')}")
        metrics.add_run(f"\n• Global Rank: {traffic_data.get('global_rank', 'N/A')}")
        metrics.add_run(f"\n• Bounce Rate: {traffic_data.get('bounce_rate', 'N/A')}")
        metrics.add_run(f"\n• Pages per Visit: {traffic_data.get('pages_per_visit', 'N/A')}")
        metrics.add_run(f"\n• Avg. Visit Duration: {traffic_data.get('avg_duration', 'N/A')}")
    
    doc.add_page_break()
    
    # TOP ORGANIC KEYWORDS
    doc.add_heading('Top Organic Keywords', level=1)
    
    top_keywords = traffic_data.get('top_keywords', []) if traffic_data else []
    if top_keywords:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Keyword'
        hdr_cells[1].text = 'Position'
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[1].paragraphs[0].runs[0].bold = True
        
        for keyword, position in top_keywords:
            row_cells = table.add_row().cells
            row_cells[0].text = keyword
            row_cells[1].text = str(position)
    else:
        doc.add_paragraph("No keyword data available.")
    
    doc.add_page_break()
    
    # TOP COUNTRIES
    doc.add_heading('Top Visitor Countries', level=1)
    
    top_countries = traffic_data.get('top_countries', []) if traffic_data else []
    if top_countries:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Country'
        hdr_cells[1].text = 'Percentage'
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[1].paragraphs[0].runs[0].bold = True
        
        for country, percentage in top_countries:
            row_cells = table.add_row().cells
            row_cells[0].text = country
            row_cells[1].text = f"{percentage}%"
    else:
        doc.add_paragraph("No country data available.")
    
    doc.add_page_break()
    
    # AI RECOMMENDATIONS
    doc.add_heading('AI Strategic Recommendations', level=1)
    
    if ai_suggestions and isinstance(ai_suggestions, list):
        for i, item in enumerate(ai_suggestions[:10], 1):
            p = doc.add_paragraph(style='List Number')
            p.add_run(f"{item.get('title', 'Recommendation')}: ").bold = True
            p.add_run(str(item.get('text', ''))[:400])
    else:
        doc.add_paragraph("No AI recommendations available.")
    
    # FOOTER
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f'Generated by {agency_name} | Powered by NexGenWebLab')
    fr.font.size = Pt(10)
    fr.font.color.rgb = SLATE
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()