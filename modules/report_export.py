import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime
import io

def create_gauge_image(score):
    color = "#ef4444" if score < 50 else "#f59e0b" if score < 90 else "#10b981"
    fig, ax = plt.subplots(figsize=(1.5, 1.5))
    ax.pie([score, max(100-score, 0)], colors=[color, '#f1f5f9'], startangle=90, counterclock=False, wedgeprops=dict(width=0.25))
    ax.text(0, 0, str(score), ha='center', va='center', fontsize=18, fontweight='bold', color='#0f172a', family='sans-serif')
    plt.axis('off')
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', transparent=True, bbox_inches='tight', dpi=150)
    plt.close(fig)
    img_stream.seek(0)
    return img_stream

def generate_word_report(url, onpage_data, speed_data, traffic_data, ai_suggestions, agency_name="SEO Agency", client_name="Client", author_name="SEO Expert"):
    doc = Document()
    doc.sections[0].page_width = Inches(8.5)
    doc.sections[0].page_height = Inches(11)
    
    PRIMARY = RGBColor(109, 40, 217)
    SUCCESS = RGBColor(16, 185, 129)
    WARNING = RGBColor(245, 158, 11)
    DANGER = RGBColor(239, 68, 68)
    SLATE = RGBColor(100, 116, 139)
    
    # COVER
    title = doc.add_heading(f'{agency_name} - SEO Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = PRIMARY
    
    sub = doc.add_heading(f'Prepared for: {client_name}', level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = SLATE
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Target URL: {url}\n").bold = True
    meta.add_run(f"Prepared by: {author_name}\n")
    meta.add_run(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
    
    doc.add_page_break()

    # PERFORMANCE SCORES
    doc.add_heading('1. Performance Scores', level=1)
    
    mobile_score = speed_data.get('mobile', {}).get('performance', 0) if speed_data else 0
    desktop_score = speed_data.get('desktop', {}).get('performance', 0) if speed_data else 0
    avg_score = (mobile_score + desktop_score) // 2
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Overall: {avg_score} | Mobile: {mobile_score} | Desktop: {desktop_score}')
    r.font.size = Pt(18)
    r.bold = True
    if avg_score >= 90:
        r.font.color.rgb = SUCCESS
    elif avg_score >= 50:
        r.font.color.rgb = WARNING
    else:
        r.font.color.rgb = DANGER
    
    doc.add_paragraph()

    # KEY METRICS
    doc.add_heading('2. Key Performance Metrics', level=1)
    
    mp = doc.add_paragraph()
    if traffic_data:
        mp.add_run(f"Global Rank: {traffic_data.get('global_rank', 'N/A')}\n")
        mp.add_run(f"Monthly Visits: {traffic_data.get('monthly_visits', 'N/A')}\n")
        mp.add_run(f"Bounce Rate: {traffic_data.get('bounce_rate', 'N/A')}\n")
        mp.add_run(f"Pages/Visit: {traffic_data.get('pages_per_visit', 'N/A')}\n")
    mp.add_run(f"Word Count: {onpage_data.get('word_count', 0)}\n")
    mp.add_run(f"Missing ALTs: {onpage_data.get('missing_alt', 0)}")
    
    doc.add_page_break()

    # ON-PAGE SEO
    doc.add_heading('3. Technical On-Page SEO', level=1)
    
    checks = [
        ('Title Tag', bool(onpage_data.get('title'))),
        ('Meta Description', bool(onpage_data.get('description'))),
        ('H1 Headings', len(onpage_data.get('h1', [])) > 0),
        ('Word Count >= 300', onpage_data.get('word_count', 0) >= 300),
        ('HTTPS/SSL', onpage_data.get('is_https', '') == 'Yes'),
        ('Schema Markup', bool(onpage_data.get('schema'))),
        ('Canonical URL', bool(onpage_data.get('canonical'))),
        ('No Missing ALTs', onpage_data.get('missing_alt', 0) == 0),
    ]
    
    for label, passed in checks:
        p = doc.add_paragraph()
        p.add_run(f"{'PASS' if passed else 'FAIL'}: {label}")
    
    doc.add_page_break()

    # CORE WEB VITALS
    doc.add_heading('4. Core Web Vitals', level=1)
    
    if speed_data:
        for device in ['mobile', 'desktop']:
            device_data = speed_data.get(device, {})
            score = device_data.get('performance', 0)
            hp = doc.add_heading(f'{device.upper()} Score: {score}/100', level=2)
            if score >= 90:
                hp.runs[0].font.color.rgb = SUCCESS
            elif score >= 50:
                hp.runs[0].font.color.rgb = WARNING
            else:
                hp.runs[0].font.color.rgb = DANGER
            
            raw_metrics = device_data.get('metrics', {})
            for key, label in [('fcp', 'FCP'), ('lcp', 'LCP'), ('tbt', 'TBT'), ('cls', 'CLS'), ('si', 'SI')]:
                value = raw_metrics.get(key, {}).get('value', 'N/A')
                doc.add_paragraph(f"  {label}: {value}")
    
    doc.add_page_break()

    # AI RECOMMENDATIONS
    doc.add_heading('5. AI Strategic Recommendations', level=1)
    
    if ai_suggestions and isinstance(ai_suggestions, list):
        for i, item in enumerate(ai_suggestions[:10], 1):
            p = doc.add_paragraph(style='List Number')
            p.add_run(f"{item.get('title', 'Recommendation')}: ").bold = True
            p.add_run(str(item.get('text', ''))[:400])
    else:
        doc.add_paragraph("No AI recommendations available.")

    # FOOTER
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f'Generated by {agency_name} | Powered by NexGenWebLab')
    fr.font.size = Pt(10)
    fr.font.color.rgb = SLATE

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()